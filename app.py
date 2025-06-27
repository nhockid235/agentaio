import streamlit as st
import tempfile
import os
import torch
import pandas as pd
from langchain_community.document_loaders import PyPDFLoader
from langchain_huggingface.embeddings import HuggingFaceEmbeddings
from langchain_experimental.text_splitter import SemanticChunker
from langchain_core.documents import Document
from langchain_chroma import Chroma
from langchain_huggingface.llms import HuggingFacePipeline
from langchain import hub
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from transformers import AutoModelForCausalLM, AutoTokenizer, pipeline

if "rag_chain" not in st.session_state:
    st.session_state.rag_chain = None
if "models_loaded" not in st.session_state:
    st.session_state.models_loaded = False
if "embeddings" not in st.session_state:
    st.session_state.embeddings = None
if "llm" not in st.session_state:
    st.session_state.llm = None

@st.cache_resource
def load_embeddings():
    return HuggingFaceEmbeddings(model_name="bkai-foundation-models/vietnamese-bi-encoder")

@st.cache_resource
def load_llm():
    MODEL_NAME = "lmsys/vicuna-7b-v1.5"
    model = AutoModelForCausalLM.from_pretrained(
        MODEL_NAME,
        torch_dtype=torch.float16,
        low_cpu_mem_usage=True
    )

    tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME, use_fast=False)
    tokenizer.pad_token = tokenizer.eos_token
    
    model_pipeline = pipeline(
        "text-generation",
        model=model,
        tokenizer=tokenizer,
        max_new_tokens=512,
        pad_token_id=tokenizer.eos_token_id,
        device_map="auto",
    )
    return HuggingFacePipeline(pipeline=model_pipeline)

def build_rag_chain_from_text(text: str):
    doc = Document(page_content=text)

    semantic_splitter = SemanticChunker(
        embeddings=st.session_state.embeddings,
        buffer_size=1,
        breakpoint_threshold_type="percentile",
        breakpoint_threshold_amount=95,
        min_chunk_size=500,
        add_start_index=True
    )
    docs = semantic_splitter.split_documents([doc])

    vector_db = Chroma.from_documents(documents=docs, embedding=st.session_state.embeddings)
    retriever = vector_db.as_retriever()

    prompt = hub.pull("rlm/rag-prompt")

    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)

    rag_chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | st.session_state.llm
        | StrOutputParser()
    )

    return rag_chain, len(docs)

def process_excel(file):
    xl = pd.ExcelFile(file)
    all_text = ""
    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name)
        all_text += df.astype(str).apply(lambda x: ' '.join(x), axis=1).str.cat(sep='\n')
    return build_rag_chain_from_text(all_text)


from docx import Document

def process_docx(file):
    from docx import Document as DocxDocument
    doc = DocxDocument(file)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return build_rag_chain_from_text(full_text)


def process_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        full_text = f.read()
    return build_rag_chain_from_text(full_text)


def process_pdf(uploaded_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_file_path = tmp_file.name

    loader = PyPDFLoader(tmp_file_path)
    documents = loader.load()

    semantic_splitter = SemanticChunker(
        embeddings=st.session_state.embeddings,
        buffer_size=1,
        breakpoint_threshold_type="percentile",
        breakpoint_threshold_amount=95,
        min_chunk_size=500,
        add_start_index=True
    )
    docs = semantic_splitter.split_documents(documents)

    vector_db = Chroma.from_documents(documents=docs, embedding=st.session_state.embeddings)
    retriever = vector_db.as_retriever()

    prompt = hub.pull("rlm/rag-prompt")

    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)

    rag_chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | st.session_state.llm
        | StrOutputParser()
    )

    return rag_chain, len(docs)

st.set_page_config(page_title="RAG Assistant", layout="wide")
st.title("🧠 Trợ lý hỏi đáp tài liệu")

st.markdown("""
**Ứng dụng AI giúp bạn hỏi đáp trực tiếp với nội dung tài liệu PDF bằng tiếng Việt**
**Cách sử dụng đơn giản:**
1. **Upload DATA** Chọn file DATA từ máy tính và nhấn "Xử lý DATA"
2. **Đặt câu hỏi** Nhập câu hỏi về nội dung tài liệu và nhận câu trả lời ngay lập tức
---
""")

if not st.session_state.models_loaded:
    st.info("Đang tải models...")
    st.session_state.embeddings = load_embeddings()
    st.session_state.llm = load_llm()
    st.session_state.models_loaded = True

uploaded_file = st.file_uploader("📁 Upload file", type=["pdf", "docx", "xlsx", "xls", "txt"])

if uploaded_file and st.button("🚀 Xử lý file"):
    with st.spinner("Đang xử lý..."):
        file_type = uploaded_file.name.split(".")[-1].lower()

        if file_type == "pdf":
            st.session_state.rag_chain, num_chunks = process_pdf(uploaded_file)
        elif file_type == "docx":
            st.session_state.rag_chain, num_chunks = process_docx(uploaded_file)
        elif file_type in ["xlsx", "xls"]:
            st.session_state.rag_chain, num_chunks = process_excel(uploaded_file)
        elif file_type == "txt":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".txt") as tmp:
                tmp.write(uploaded_file.getvalue())
                tmp_path = tmp.name
            st.session_state.rag_chain, num_chunks = process_txt(tmp_path)
        else:
            st.error("❌ Định dạng file không được hỗ trợ!")
            st.stop()

        st.success(f"✅ Đã xử lý xong! Tổng cộng: {num_chunks} chunks")


if st.session_state.rag_chain:
    question = st.text_input("💬 Đặt câu hỏi về nội dung tài liệu:")
    if question:
        with st.spinner("🤖 Đang suy nghĩ..."):
            output = st.session_state.rag_chain.invoke(question)
            answer = output.split("Answer:")[1].strip() if "Answer:" in output else output.strip()
            st.markdown("### 📝 **Trả lời:**")
            st.write(answer)
